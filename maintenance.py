import datetime
import getopt
import math
import random
import sys
import time

from docx import Document


def getParameters(argv):
    number_of_equations, list_of_types, title = 10,'u','Default'
    try:
        opts, args = getopt.getopt(argv,"hn:t:T:",["help","number=","types=","Title="])
    except getopt.GetoptError:
        print("use aufgaben -n 4 -t usn -T 'Kurztest 2'")
        sys.exit(2)
    for opt, arg in opts:
        if opt =='-h':
            print("use exerc_equations -n 4 -t usn -T 'Kurztest 2'")
            sys.exit()
        elif opt in ("-n","--number"):
            number_of_equations = int(arg)
        elif opt in ("-t","--types"):
            list_of_types = list(arg)
        elif opt in ("-T","--Title"):
            title = arg
        else:
            sys.exit()
        
    return [number_of_equations, list_of_types, title]


def check_types(t: dict, lot: list)-> list:
    for item in lot:
        if item not in t:
            lot.remove(item)
    return lot

def createFile(funs, name ="Aufgaben", intro = "Solve"):
    """
    creates word document , exercises is a list of (equation, solution) pairs
    """
    #weeknumber = datetime.datetime.fromtimestamp(time.time()).isocalendar()[1]
    timestmp = datetime.date.fromtimestamp(time.time())
    document = Document()
    
    document.core_properties.author="Gregor Lüdi"
    document.core_properties.title=name
    document.core_properties.subject="Übungsaufgaben"
    document.core_properties.created = datetime.datetime.today()
    document.core_properties.keywords = "Informatik, Übungaufgaben"
    document.core_properties.comments = "Automatisch generiertes Dokument "
    document.add_heading('Übungsaufgaben', 0)
    document.add_paragraph(intro)

    document.add_heading('Aufgaben', level=1)
    document.add_paragraph(f"{funs[0][0]}")
    for item in funs:
        document.add_paragraph(f"{item[1]},", style='List Number')

    document.add_page_break()

    document.add_heading('Lösungen', level=1)
    for item in funs:
        document.add_paragraph("{} ".format(item[2]),style='List Number 2')

    document.save('{}_{}.docx'.format(name, timestmp))


